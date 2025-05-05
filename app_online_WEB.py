import os
import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
import platform
import streamlit.components.v1 as components

# === Preparação de pastas ===
os.makedirs("laudos_exportados", exist_ok=True)
os.makedirs("backups", exist_ok=True)

# === Carregamento da base de frases ===
df = pd.read_csv("banco_frases_app.csv")
alteracoes = df["nome_da_alteracao"].unique()

# === Seleção do tipo de exame ===
st.title("Painel Lili - Versão Online (com Voz via Navegador)")
tipo_exame = st.selectbox("Escolha o tipo de exame:", ["abdome", "tireoide", "mama", "musculoesquelético"])
template_path = f"templates/{tipo_exame}.docx"
doc = Document(template_path)

# === Entradas de cabeçalho ===
st.subheader("Informações do Paciente")
nome_paciente = st.text_input("Nome do paciente")
data_nascimento = st.text_input("Data de nascimento")
medico_solicitante = st.text_input("Médico solicitante")

# === Campo de alteração ===
st.subheader("Frases clínicas")
col1, col2 = st.columns(2)
with col1:
    alteracao_escolhida = st.selectbox("Nome da alteração clínica", sorted(alteracoes))
with col2:
    texto_alteracao = st.text_area("Descrição da alteração", value="")

if alteracao_escolhida:
    frase = df[df["nome_da_alteracao"] == alteracao_escolhida].iloc[0]
    texto_alteracao = frase["descricao_da_alteracao"]
    conclusao_automatica = frase["conclusao_da_alteracao"]
else:
    texto_alteracao = ""
    conclusao_automatica = ""

# === Campo de edição completa do laudo ===
st.subheader("Pré-visualização e edição do laudo")
preview_text = st.text_area("Laudo completo (editável)", height=400, key="preview_text")

# === Integração com Web Speech API ===
st.subheader("Comando de Voz (navegador)")

components.html("""
    <script>
    const streamlitReceive = (text) => {
        const input = window.parent.document.querySelector('textarea[data-baseweb="textarea"]');
        if (input) {
            input.value += '\n' + text;
            const event = new Event('input', { bubbles: true });
            input.dispatchEvent(event);
        }
    }

    function iniciarReconhecimento() {
        var recognition = new webkitSpeechRecognition() || new SpeechRecognition();
        recognition.lang = 'pt-BR';
        recognition.interimResults = false;
        recognition.maxAlternatives = 1;

        recognition.start();
        recognition.onresult = function(event) {
            var comando = event.results[0][0].transcript;
            streamlitReceive(comando);
        };
        recognition.onerror = function(event) {
            console.error('Erro:', event.error);
        };
    }
    </script>
    <button onclick="iniciarReconhecimento()">🎙️ Falar</button>
""", height=100)

# === Salvar o documento ===
if st.button("Finalizar e Salvar Laudo"):
    doc.add_paragraph(f"Paciente: {nome_paciente}")
    doc.add_paragraph(f"Data de nascimento: {data_nascimento}")
    doc.add_paragraph(f"Médico solicitante: {medico_solicitante}")
    doc.add_paragraph("")  # espaçamento
    doc.add_paragraph(st.session_state["preview_text"])
    doc.add_paragraph("")
    doc.add_paragraph("Conclusão:")
    doc.add_paragraph(conclusao_automatica)
    doc.add_paragraph("")
    doc.add_paragraph("Dr. Andrew Costa | CRM 71703/MG")

    data_str = datetime.now().strftime("%d-%m-%Y")
    nome_formatado = nome_paciente.replace(" ", "_")
    tipo_formatado = tipo_exame.lower().replace("musculoesquelético", "msk").replace(" ", "_")
    nome_arquivo = f"{data_str}_{nome_formatado}_{tipo_formatado}.docx"
    caminho = os.path.join("laudos_exportados", nome_arquivo)
    doc.save(caminho)

    sistema = platform.system()
    if sistema == "Darwin":
        os.system(f'open "{caminho}"')
    elif sistema == "Windows":
        os.startfile(caminho)

    st.success(f"Laudo salvo em: {caminho}")