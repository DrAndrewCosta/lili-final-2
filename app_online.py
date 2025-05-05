import os
import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
import platform
from vosk import Model, KaldiRecognizer
import sounddevice as sd
import queue
import json
import tempfile

# === Preparacao de pastas ===
os.makedirs("laudos_exportados", exist_ok=True)
os.makedirs("backups", exist_ok=True)

# === Carregamento da base de frases ===
df = pd.read_csv("banco_frases_app.csv")
alteracoes = df["nome_da_alteracao"].unique()

# === Selecao do tipo de exame ===
st.title("Painel Lili - Geracao de Laudo com Voz e Texto")
tipo_exame = st.selectbox("Escolha o tipo de exame:", ["abdome", "tireoide", "mama", "musculoesqueletico"])
template_path = f"templates/{tipo_exame}.docx"
doc = Document(template_path)

# === Entradas de cabecalho ===
st.subheader("Informacoes do Paciente")
nome_paciente = st.text_input("Nome do paciente")
data_nascimento = st.text_input("Data de nascimento")
medico_solicitante = st.text_input("Medico solicitante")

# === Campo para escolha da alteracao ===
st.subheader("Frases clinicas")
col1, col2 = st.columns(2)
with col1:
    alteracao_escolhida = st.selectbox("Nome da alteracao clinica", sorted(alteracoes))
with col2:
    texto_alteracao = st.text_area("Descricao da alteracao", value="")

# === Atualizacao automatica do texto ao escolher alteracao ===
if alteracao_escolhida:
    frase = df[df["nome_da_alteracao"] == alteracao_escolhida].iloc[0]
    texto_alteracao = frase["descricao_da_alteracao"]
    conclusao_automatica = frase["conclusao_da_alteracao"]
else:
    texto_alteracao = ""
    conclusao_automatica = ""

# === Campo de edicao manual do laudo ===
st.subheader("Pre-visualizacao e edicao do laudo")
preview_text = st.text_area("Laudo completo (editavel)", height=400)

# === Funcao para escuta por voz ===
st.subheader("Comando por Voz")
if "voz_lili" not in st.session_state:
    st.session_state.voz_lili = []

def escutar_comando():
    model = Model("voz_lili/vosk-model-small-pt-0.3")
    rec = KaldiRecognizer(model, 16000)
    q = queue.Queue()

    def callback(indata, frames, time, status):
        q.put(bytes(indata))

    with sd.RawInputStream(samplerate=16000, blocksize=8000, dtype='int16',
                           channels=1, callback=callback):
        while True:
            data = q.get()
            if rec.AcceptWaveform(data):
                result = rec.Result()
                texto = json.loads(result)["text"]
                return texto

# === Salvar o documento ===
if st.button("Finalizar e Salvar Laudo"):
    doc.add_paragraph(f"Paciente: {nome_paciente}")
    doc.add_paragraph(f"Data de nascimento: {data_nascimento}")
    doc.add_paragraph(f"Medico solicitante: {medico_solicitante}")
    doc.add_paragraph("")  # espacamento
    doc.add_paragraph(preview_text)
    doc.add_paragraph("")
    doc.add_paragraph("Conclusao:")
    doc.add_paragraph(conclusao_automatica)
    doc.add_paragraph("")
    doc.add_paragraph("Dr. Andrew Costa | CRM 71703/MG")

    data_str = datetime.now().strftime("%d-%m-%Y")
    nome_formatado = nome_paciente.replace(" ", "_")
    tipo_formatado = tipo_exame.lower().replace("musculoesqueletico", "msk").replace(" ", "_")
    nome_arquivo = f"{data_str}_{nome_formatado}_{tipo_formatado}.docx"
    caminho = os.path.join("laudos_exportados", nome_arquivo)
    doc.save(caminho)

    sistema = platform.system()
    if sistema == "Darwin":
        os.system(f'open "{caminho}"')
    elif sistema == "Windows":
        os.startfile(caminho)

    st.success(f"Laudo salvo em: {caminho}")